"""
services/word_generator.py

Генератор .docx из списка блоков JSON.
Типы блоков (как на фронте): title-page, heading, text, image, table

Структура документа:
  Страница 1 — титульная (title-page)
  Страница 2 — Содержание (динамическое, обновляется Word автоматически)
  Страницы 3+ — контент, H1 = новая страница
"""

import io
import logging
import requests

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
# Стили по ГОСТ 7.32
# ─────────────────────────────────────────────

def _setup_gost_styles(doc: Document):
    """Настройка глобальных стилей документа."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.25)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for i in range(1, 4):
        h_style = doc.styles[f"Heading {i}"]
        h_style.font.name = "Times New Roman"
        h_style.font.bold = True
        h_style.font.color.rgb = RGBColor(0, 0, 0)
        h_style.font.size = Pt(16 if i == 1 else 14)
        h_pf = h_style.paragraph_format
        h_pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_pf.first_line_indent = Cm(1.25)
        h_pf.space_before = Pt(12)
        h_pf.space_after = Pt(6)
        h_pf.keep_with_next = True


# ─────────────────────────────────────────────
# Утилиты
# ─────────────────────────────────────────────

def _set_font(run, size_pt: int, bold: bool = False, italic: bool = False,
              font_name: str = "Times New Roman"):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic


def _fetch_image_bytes(url: str):
    try:
        resp = requests.get(url, timeout=10)
        resp.raise_for_status()
    except Exception as e:
        logger.warning("Не удалось загрузить изображение %s: %s", url, e)
        return None

    content_type = resp.headers.get("Content-Type", "")
    if "svg" in content_type or url.lower().endswith(".svg"):
        logger.info("SVG пропущен (не поддерживается Word): %s", url)
        return None

    return resp.content, content_type


def _page_break_into(p):
    """
    Вставляет разрыв страницы как run внутрь существующего параграфа p.
    Не создаёт лишний пустой параграф в начале следующей страницы,
    в отличие от doc.add_page_break().
    """
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _add_page_number(doc: Document):
    """
    Номер страницы в нижнем колонтитуле со 2-й страницы.
    Титульная (1-я) — без колонтитула.
    """
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Обычный футер (2-я страница и далее)
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run()

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

    # Футер первой страницы — явно пустой
    first_footer = section.first_page_footer
    if first_footer.paragraphs:
        first_footer.paragraphs[0].clear()


# ─────────────────────────────────────────────
# Содержание (динамическое TOC)
# ─────────────────────────────────────────────

def _render_dynamic_toc(doc: Document):
    """
    Вставляет TOC-поле.
    - Microsoft Word: обновится автоматически при открытии (updateFields в XML).
    - LibreOffice: нужно вручную нажать F9 или "Обновить всё".
    """
    # Заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("СОДЕРЖАНИЕ")
    _set_font(run, 16, bold=True)

    # Параграф с полем TOC
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Cm(0)
    run = paragraph.add_run()

    # begin
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    fldChar1.set(qn("w:dirty"), "true")   # пометить как «нужно обновить»
    run._r.append(fldChar1)

    # инструкция TOC
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)

    # separate — между инструкцией и placeholder-текстом
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)

    # placeholder — без него некоторые версии Word не рендерят TOC
    placeholder = OxmlElement("w:r")
    placeholder_rpr = OxmlElement("w:rPr")
    placeholder_noProof = OxmlElement("w:noProof")
    placeholder_rpr.append(placeholder_noProof)
    placeholder.append(placeholder_rpr)
    placeholder_t = OxmlElement("w:t")
    placeholder_t.text = "Обновите содержание: нажмите правой кнопкой → Обновить поле"
    placeholder.append(placeholder_t)
    paragraph._p.append(placeholder)

    # end + разрыв страницы внутри того же параграфа
    run2 = paragraph.add_run()
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run2._r.append(fldChar3)
    _page_break_into(paragraph)

    # Принудительное updateFields через XML settings (надёжнее чем .update_fields)
    settings_element = doc.settings.element
    updateFields = OxmlElement("w:updateFields")
    updateFields.set(qn("w:val"), "true")
    settings_element.append(updateFields)


# ─────────────────────────────────────────────
# Рендереры блоков
# ─────────────────────────────────────────────

from docx.enum.table import WD_ALIGN_VERTICAL

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.shared import Cm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _render_title_page(doc: Document, content: dict):
    # Динамические данные
    university = content.get("university", "")
    department = content.get("department", "__________________________________________")
    subject    = content.get("subject", "")
    title      = content.get("title", "")
    student    = content.get("studentName", "")
    group      = content.get("group", "")
    teacher    = content.get("teacherName", "")
    rank       = content.get("teacherRank", "") # канд. экон. наук, доцент
    city       = content.get("city", "МИНСК")
    year       = content.get("year", "2026")

    # Основная таблица-каркас (3 строки: верх, центр, низ)
    table = doc.add_table(rows=3, cols=1)
    table.autofit = False
    
    # Убираем границы
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for b in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{b}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tbl_pr.append(borders)

    # =========================
    # 🔹 ВЕРХ (ВУЗ)
    # =========================
    top = table.cell(0, 0)
    top.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    
    p = top.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    _set_font(p.add_run(university.upper()), 12, bold=True)

    p2 = top.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(20)
    _set_font(p2.add_run(f"Кафедра {department}"), 12)

    # =========================
    # 🔹 ЦЕНТР (ТЕМА И ПОДПИСИ)
    # =========================
    mid = table.cell(1, 0)
    mid.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    # Заголовок работы
    p_work = mid.paragraphs[0]
    p_work.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p_work.add_run("КУРСОВАЯ РАБОТА\n"), 16, bold=True)

    p_title = mid.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Убрали явное "по дисциплине", просто выводим значения
    _set_font(p_title.add_run(f"{subject}\n"), 14)
    _set_font(p_title.add_run(f"{title}"), 14, bold=True)

    # Блок подписей (вложенная таблица для выравнивания "по линии")
    mid.add_paragraph().paragraph_format.space_before = Pt(50)
    # Создаем таблицу в правой части (2 строки, 3 колонки)
    # Колонки: 1 - Должность, 2 - Подписи (сдвинутые вниз), 3 - ФИО
    inner = mid.add_table(rows=2, cols=3)
    inner.alignment = WD_ALIGN_PARAGRAPH.RIGHT # Сдвигаем весь блок вправо
    
    # Настройка ширины (примерно) для ровных вертикальных линий
    inner.columns[0].width = Cm(4.5)
    inner.columns[1].width = Cm(4.0)
    inner.columns[2].width = Cm(4.5)

    # --- СТУДЕНТ ---
    row1 = inner.rows[0].cells
    # Ячейка 1: Студент (без отступов)
    p_st = row1[0].paragraphs[0]
    _set_font(p_st.add_run(f"Студент\n{group}"), 12)
    
    # Ячейка 2: Подпись/Дата (сдвигаем ниже для места под ручку)
    p_sig1 = row1[1].paragraphs[0]
    p_sig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig1.paragraph_format.space_before = Pt(18) # Сдвиг вниз
    _set_font(p_sig1.add_run("(подпись)\n(дата)"), 8)
    
    # Ячейка 3: ФИО
    p_name1 = row1[2].paragraphs[0]
    p_name1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(p_name1.add_run(student), 12)

    # --- РУКОВОДИТЕЛЬ ---
    row2 = inner.rows[1].cells
    # Ячейка 1: Должность
    p_rep = row2[0].paragraphs[0]
    _set_font(p_rep.add_run(f"Руководитель\n{rank}"), 12)
    
    # Ячейка 2: Оценка/Подпись (сдвигаем ниже)
    p_sig2 = row2[1].paragraphs[0]
    p_sig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sig2.paragraph_format.space_before = Pt(18) # Сдвиг вниз
    _set_font(p_sig2.add_run("(подпись) (оценка)\n(дата)"), 8)
    
    # Ячейка 3: ФИО
    p_name2 = row2[2].paragraphs[0]
    p_name2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_font(p_name2.add_run(teacher), 12)

    # =========================
    # 🔹 НИЗ (ГОРОД)
    # =========================
    bot = table.cell(2, 0)
    bot.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    
    p_bot = bot.paragraphs[0]
    p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_font(p_bot.add_run(f"{city.upper()} {year}"), 12)

    # Фиксация высот
    table.rows[0].height = Cm(4)
    table.rows[1].height = Cm(17.5) 
    table.rows[2].height = Cm(2)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def _render_heading(doc: Document, content: dict):
    """heading: { level: 1|2|3, text: "..." }
    Используем встроенный стиль Heading N — заголовки попадают в динамический TOC.
    H1 по ГОСТ — заглавными буквами.
    """
    level = content.get("level", 1)
    text  = content.get("text", "")
    if level == 1:
        text = text.upper()
    p = doc.add_paragraph(text, style=f"Heading {level}")
    # Явно форсируем Times New Roman — стиль Heading может наследовать другой шрифт темы
    for run in p.runs:
        run.font.name = "Times New Roman"


def _render_text(doc: Document, content: dict):
    """text: { text: "..." } — стиль Normal (полуторный, отступ 1.25, по ширине)."""
    raw = content.get("text", "")
    for line in raw.split("\n"):
        if line.strip():
            doc.add_paragraph(line)


def _render_image(doc: Document, content: dict, img_counter: list):
    """image: { url: "...", caption: "..." }
    Подпись: "Рисунок N — caption" (формат ГОСТ).
    """
    url     = content.get("url", "")
    caption = content.get("caption", "")
    img_counter[0] += 1
    num = img_counter[0]

    if not url:
        p = doc.add_paragraph(f"[Изображение {num} не указано]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return

    result = _fetch_image_bytes(url)
    if result is None:
        p = doc.add_paragraph(f"[Изображение {num} недоступно или SVG: {url}]")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        img_bytes, _ = result
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            run = p.add_run()
            run.add_picture(io.BytesIO(img_bytes), width=Cm(14))
        except Exception as e:
            logger.warning("Ошибка вставки изображения %s: %s", url, e)
            p = doc.add_paragraph(f"[Ошибка вставки изображения {num}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if caption:
        cap_p = doc.add_paragraph(f"Рисунок {num} — {caption}")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.first_line_indent = Cm(0)
        _set_font(cap_p.runs[0], 12, italic=True)


def _render_table(doc: Document, content: dict):
    """table: { rows: N, cols: M, data: ["cell00", "cell01", ...] }
    Первая строка — заголовок (bold + серый фон). Шрифт 12pt по ГОСТ.
    """
    rows_count = content.get("rows", 1)
    cols_count = content.get("cols", 1)
    data       = content.get("data", [])

    table = doc.add_table(rows=rows_count, cols=cols_count)
    table.style = "Table Grid"

    for r in range(rows_count):
        for c in range(cols_count):
            cell = table.cell(r, c)
            idx  = r * cols_count + c
            cell.text = str(data[idx]) if idx < len(data) else ""

            for para in cell.paragraphs:
                para.paragraph_format.first_line_indent = Cm(0)
                for run in para.runs:
                    run.font.size = Pt(12)
                    if r == 0:
                        run.bold = True

            if r == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:color"), "auto")
                shading.set(qn("w:fill"), "D9D9D9")
                cell._tc.get_or_add_tcPr().append(shading)

    # Отступ после таблицы через space_after на последней строке, без пустого параграфа
    for cell in table.rows[-1].cells:
        for para in cell.paragraphs:
            para.paragraph_format.space_after = Pt(12)


# ─────────────────────────────────────────────
# Точка входа
# ─────────────────────────────────────────────

def generate_docx_from_blocks(blocks: list) -> Document:
    doc = Document()
    _setup_gost_styles(doc)

    # Поля по ГОСТ: левое 30мм, правое 15мм, верх/низ 20мм
    section = doc.sections[0]
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    _add_page_number(doc)

    # ── Страница 1: Титульная ──
    title_block = next((b for b in blocks if b.get("type") == "title-page"), None)
    if title_block:
        _render_title_page(doc, title_block.get("content", {}))
        # Переносим разрыв СЮДА, после таблицы, чтобы не ломать высоту строк
        # doc.add_page_break()

    # ── Страница 2: Содержание ──
    _render_dynamic_toc(doc)

    # ── Страницы 3+: Контент ──
    # ... (остальной код без изменений)
    content_blocks = [b for b in blocks if b.get("type") != "title-page"]
    img_counter = [0]
    last_paragraph = [None]

    for i, block in enumerate(content_blocks):
        b_type  = block.get("type")
        content = block.get("content", {})

        if b_type == "heading" and content.get("level") == 1 and i > 0:
            if last_paragraph[0] is not None:
                _page_break_into(last_paragraph[0])
            else:
                p = doc.add_paragraph()
                _page_break_into(p)

        try:
            if b_type == "heading":
                _render_heading(doc, content)
                last_paragraph[0] = doc.paragraphs[-1]
            elif b_type == "text":
                _render_text(doc, content)
                last_paragraph[0] = doc.paragraphs[-1]
            elif b_type == "image":
                _render_image(doc, content, img_counter)
                last_paragraph[0] = doc.paragraphs[-1]
            elif b_type == "table":
                _render_table(doc, content)
                last_paragraph[0] = doc.paragraphs[-1]
        except Exception as e:
            logger.error("Ошибка при рендере блока %s: %s", b_type, e)

    return doc